from datetime import datetime, timedelta

from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from numpy import random


class quiz_list_generator(object):
    def __init__(self, element_a=99, element_b=99, operations="+-",
                 min_result=0, max_result=99, num_of_quizes=20, if_with_blanks="1000"):
        self.element_a = element_a
        self.element_b = element_b
        self.min_result = min_result
        self.max_result = max_result
        self.num_of_quizes = num_of_quizes
        self.if_with_blanks = (if_with_blanks + "000")[:4]
        self.operations = []
        for i in range(len(operations)):
            if "+-*/".find(operations[i]) >= 0 and operations[i] not in self.operations:
                self.operations.append(operations[i])
        self.quiz_per_page = []

    def generate_quizes(self):

        def get_qz(a, op, b):
            return a + " " + op + " " + b + " " + "="

        quiz_array = []
        flag = 0
        while True:
            a = str(random.randint(self.element_a))
            if not flag:
                op = random.choice(self.operations)
            else:
                op = "/"
            b = str(random.randint(self.element_b))

            qz = get_qz(a, op, b)

            # prevent the zero division
            try:
                rst = eval(qz[:-1])
            except ZeroDivisionError:
                rst = 0.0001

            # prevent the float
            if rst != int(rst):
                flag = 1
            elif self.min_result <= rst <= self.max_result:
                if self.if_with_blanks[self.operations.index(op)] != "0" and random.randint(100) <= 50:
                    if random.randint(100) < 50:
                        qz = get_qz("__", op, b)
                        quiz_array.append(qz + " " + str(rst))
                    else:
                        qz = get_qz(a, op, "__")
                        quiz_array.append(qz + " " + str(rst))
                else:
                    quiz_array.append(qz + " " + "___")

                flag = 0
            if len(quiz_array) >= self.num_of_quizes and not flag:
                break

        for i in range(0, len(quiz_array) - 1, 2):
            self.quiz_per_page.append([quiz_array[i], quiz_array[i + 1]])


def generate_one_day(date, page_per_day, file_name_surfix):
    quiz_doc_file = []
    for i in range(page_per_day):
        qz = quiz_list_generator(element_a=ELEMENT_A, element_b=ELEMENT_B, operations=OPERATIONS,
                                 min_result=MIN_RESULT, max_result=MAX_RESULT, num_of_quizes=NUM_OF_QUIZES,
                                 if_with_blanks=IF_WITH_BLANKS)
        quiz_doc_file.append(qz)
        quiz_doc_file[i].generate_quizes()
    create_doc(date, page_per_day, file_name_surfix, quiz_doc_file)


def create_doc(date, page_per_day, file_name_surfix, quiz_doc_file):
    global quiz_font_size
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = '仿宋'
    font.size = Pt(10)

    title1 = date + ' -- '
    title3 = ' of integers between '

    for page in range(page_per_day):
        title2 = OPERATIONS
        title4 = str(MIN_RESULT) + "~" + str(MAX_RESULT)
        # title creation
        p = document.add_heading(level=0)
        wp = p.add_run(title1)
        wp.font.size = Pt(15)
        wp.font.bold = True
        wp.font.color.rgb = RGBColor(0, 0, 0)

        wp = p.add_run(title2)
        wp.font.size = Pt(15)
        wp.font.color.rgb = RGBColor(255, 0, 0)

        wp = p.add_run(title3)
        wp.font.size = Pt(15)
        wp.font.bold = True
        wp.font.color.rgb = RGBColor(0, 0, 0)

        wp = p.add_run(title4 + '\n')
        wp.font.size = Pt(15)
        wp.font.color.rgb = RGBColor(255, 0, 0)

        wp = p.add_run('整数 ')
        wp.font.size = Pt(15)
        wp.font.bold = True
        wp.font.color.rgb = RGBColor(0, 0, 0)

        wp = p.add_run(title4)
        wp.font.size = Pt(15)
        wp.font.color.rgb = RGBColor(255, 0, 0)

        wp = p.add_run(' 之间的')
        wp.font.size = Pt(15)
        wp.font.bold = True
        wp.font.color.rgb = RGBColor(0, 0, 0)

        wp = p.add_run(title2)
        wp.font.size = Pt(15)
        wp.font.color.rgb = RGBColor(255, 0, 0)

        wp = p.add_run('练习 -- 第 ' + str(page + 1) + ' 页')
        wp.font.size = Pt(15)
        wp.font.bold = True
        wp.font.color.rgb = RGBColor(0, 0, 0)

        # print(quiz_doc_file[page].type_per_page)
        # print(quiz_doc_file[page].quiz_per_page)
        quiz_font_size = int(22 - (NUM_OF_QUIZES - 22) * 3)
        if quiz_font_size >= 30:
            quiz_font_size = 30
        print(quiz_font_size)

        masked_table = quiz_doc_file[page].quiz_per_page
        print(masked_table)
        table = document.add_table(rows=1, cols=2)

        for q1, q2 in masked_table:
            row_cells = table.add_row().cells
            row_cells[0].text = q1
            row_cells[1].text = q2

        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size = Pt(quiz_font_size)
        if page < page_per_day - 1:
            document.add_page_break()
    try:
        document.save('f:\\___SOPHIA____\\--Sophia K3 Folder\\MATH_PRACTICE\\math_quiz' + file_name_surfix + '_' + \
                      date + '.docx')
    except FileNotFoundError:
        print('ERROR')
        document.save('temp.docx')


def get_genral_input():
    # get the dates to print
    try:
        adv_date = int(input('How many days from now? (default = 1(today), max = 30days)')) + 0
        if adv_date > 30 or adv_date <= 0:
            adv_date = 1
    except ValueError:
        adv_date = 1

    # get how many pages per day:
    try:
        page_per_day = int(input('How many pages per day? (default = 3 pages, max = 5 pages)'))
        if page_per_day >5 or page_per_day <= 0:
            page_per_day = 1
    except ValueError:
        page_per_day = 3

    try:
        # get max quizes per page:
        max_quiz_per_page = int(input('Pleas input max quiz per page? (default = 22 quiz)')) // 2 * 2
        if max_quiz_per_page > 22 or max_quiz_per_page < 10:
            max_quiz_per_page = 22
    except ValueError:
        max_quiz_per_page = 22
    print(adv_date, page_per_day, (max_quiz_per_page // 2) * 2)

    # get file_name_surfix
    file_name_surfix = input("file name surfix:")
    if file_name_surfix != '':
        file_name_surfix = '_' + file_name_surfix

    return adv_date, page_per_day, file_name_surfix, max_quiz_per_page


def main():
    global quiz_font_size, NUM_OF_QUIZES
    quiz_font_size = 22
    adv_date, page_per_day, file_name_surfix, max_quiz_per_page = get_genral_input()
    NUM_OF_QUIZES = max_quiz_per_page
    print(NUM_OF_QUIZES)
    for i in range(adv_date):
        date = (datetime.now() + timedelta(days=i)).date().strftime('%Y_%m_%d')
        generate_one_day(date, page_per_day, file_name_surfix)
    pass


# ELEMENT_A = 10
# ELEMENT_B = 5
# OPERATIONS = "*"
# MIN_RESULT = 1
# MAX_RESULT = 99
# NUM_OF_QUIZES = 20
# IF_WITH_BLANKS = "0"

ELEMENT_A = 99
ELEMENT_B = 99
OPERATIONS = "+-"
MIN_RESULT = 0
MAX_RESULT = 99
NUM_OF_QUIZES = 20
IF_WITH_BLANKS = "10"

if __name__ == '__main__':
    main()

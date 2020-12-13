import random
from datetime import datetime, timedelta

from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor


def main():
    tables = []
    for pattern in patterns:
        table = get_table(pattern)
        tables.append(table)
    print(tables)
    file_name_surfix = 'addd-multi'
    date = (datetime.now() + timedelta(days=0)).date().strftime('%Y_%m_%d')
    create_doc(date, len(patterns), file_name_surfix, tables)

def get_single_quiz(number, times):
    temp = ''
    for i in range(times):
        if i != 0:
            temp += '+'
        temp += str(number)
    temp += ' = '
    return temp


def get_table(pattern):
    number = random.randint(pattern[0], pattern[1])
    table = []
    for i in range(12):
        table.append(get_single_quiz(number, i + 1))
    if pattern[2]:
        random.shuffle(table)
    return table


def create_doc(date, page_per_day, file_name_surfix, tables):
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = '仿宋'
    font.size = Pt(10)

    title1 = date + ' -- '
    title3 = ' of integers between '

    for page in range(page_per_day):
        title2 = ' additons to muliplications '
        title5 = '加法转乘法'
        title4 = '1-12'
        # title creation
        p = document.add_heading(level=0)
        wp = p.add_run(title1)
        wp.font.size = Pt(15)
        wp.font.bold = True
        wp.font.color.rgb = RGBColor(0, 0, 0)

        wp = p.add_run(title2)
        wp.font.size = Pt(15)
        wp.font.color.rgb = RGBColor(255, 0, 0)

        wp = p.add_run(title3)
        wp.font.size = Pt(15)
        wp.font.bold = True
        wp.font.color.rgb = RGBColor(0, 0, 0)

        wp = p.add_run(title4 + '\n')
        wp.font.size = Pt(15)
        wp.font.color.rgb = RGBColor(255, 0, 0)

        wp = p.add_run('整数 ')
        wp.font.size = Pt(15)
        wp.font.bold = True
        wp.font.color.rgb = RGBColor(0, 0, 0)

        wp = p.add_run(title4)
        wp.font.size = Pt(15)
        wp.font.color.rgb = RGBColor(255, 0, 0)

        wp = p.add_run(' 之间的')
        wp.font.size = Pt(15)
        wp.font.bold = True
        wp.font.color.rgb = RGBColor(0, 0, 0)

        wp = p.add_run(title5)
        wp.font.size = Pt(15)
        wp.font.color.rgb = RGBColor(255, 0, 0)

        wp = p.add_run('练习 -- 第 ' + str(page + 1) + ' 页')
        wp.font.size = Pt(15)
        wp.font.bold = True
        wp.font.color.rgb = RGBColor(0, 0, 0)

        quiz_font_size = 20

        table = document.add_table(rows=1, cols=1)

        for q1 in tables[page]:
            row_cells = table.add_row().cells
            row_cells[0].text = q1


        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size = Pt(quiz_font_size)
        if page < page_per_day - 1:
            document.add_page_break()
    try:
        document.save('f:\\___SOPHIA____\\--Sophia K2 Folder\\MATH_PRACTICE\\math_quiz' + file_name_surfix + '_' + \
                      date + '.docx')
    except FileNotFoundError:
        document.save('temp.docx')


if __name__ == '__main__':
    patterns = []
    shuffle = 0
    for i in range(5):
        patterns.append([i + 1, i + 1, shuffle])

    main()


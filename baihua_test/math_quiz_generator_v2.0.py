import random
import re
from datetime import datetime, timedelta

from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor


# translation function
def trans(calc):
    tmp = ''
    if calc == 'division':
        tmp = '除法'
    elif calc == 'subtraction':
        tmp = '减法'
    elif calc == 'multiplication':
        tmp = '乘法'
    else:
        tmp = '加法'
    return tmp


# create quiz list
def quiz_create(field, obj, calc, quiz_number, max_result, elements):
    calc_dic = {'addition': '+',
                'subtraction': '-',
                'multiplication': 'x',
                'division': '÷'}
    length = quiz_number // 2 * 2
    subjs = []
    for _ in range(0, length):
        subjs.append(random.randint(field[0], field[1]))
    quizs = []
    # more than 2 elements calcs
    if elements == 3:
        for subj in subjs:
            num_a = subj
            num_b = random.randint(field[0], field[1])
            while True:
                tmp = [num_a, num_b, random.choice(obj)]
                chk_result = eval(str(tmp[0]) + calc_dic[calc] + str(tmp[1])+ calc_dic[calc] + str(tmp[2]))
                if tmp not in quizs and 0 <= chk_result <= max_result:
                    #and [tmp[1], tmp[0]] not in quizs and [tmp[0], tmp[1]] not in quizs \
                    break
                else:
                    num_a = random.randint(field[0], field[1])
                    num_b = random.randint(field[0], field[1])

            random.shuffle(tmp)
            quizs.append(tmp)
        quiz_table = []
        quiz_table1 = []
        for i in range(0, len(quizs), 2):
            quiz_table.append((str(quizs[i][0]) + ' ' + calc_dic[calc] + ' ' + str(quizs[i][1]) +  ' ' + calc_dic[calc]
                               + ' ' + str(quizs[i][2]) + ' =',
                               str(quizs[i + 1][0]) + ' ' + calc_dic[calc] + ' ' + str(quizs[i + 1][1]) + ' '
                               + calc_dic[calc] + ' ' + str(quizs[i + 1][2]) + ' ='))
        ready_quiz_table = make_up_final_quiz(quiz_table + quiz_table1)
    else:
        # just 2 elements calcs
        for subj in subjs:
            num_a = subj
            while True:
                tmp = [num_a, random.choice(obj)]
                chk_result = eval(str(tmp[0]) + calc_dic[calc] + str(tmp[1]))
                if tmp not in quizs and 0 <= chk_result <= max_result:
                    # and [tmp[1], tmp[0]] not in quizs and [tmp[0], tmp[1]] not in quizs \
                    break
                else:
                    num_a = random.randint(field[0], field[1])

            random.shuffle(tmp)
            quizs.append(tmp)
        quiz_table = []
        quiz_table1 = []
        for i in range(0, len(quizs), 2):
            quiz_table.append((str(quizs[i][0]) + ' ' + calc_dic[calc] + ' ' + str(quizs[i][1]) + ' =',
                               str(quizs[i + 1][0]) + ' ' + calc_dic[calc] + ' ' + str(quizs[i + 1][1]) + ' ='))
        ready_quiz_table = make_up_final_quiz(quiz_table + quiz_table1)

    return ready_quiz_table


def make_up_final_quiz(table):
    new_table = []
    # put actual result to the table
    for quiz_tuple in table:
        new_table.append((quiz_tuple[0] + ' ' + str(eval(quiz_tuple[0].replace('=', ''))),
                          quiz_tuple[1] + ' ' + str(eval(quiz_tuple[1].replace('=', '')))))
    return new_table


def mask_table(table, mask, elements):
    new_table = []
    calc_dic = {'addition': '+',
                'subtraction': '-',
                'multiplication': 'x',
                'division': '÷'}
    for quiz_tuple in table:
        if mask == 0:
            new_table.append((quiz_tuple[0][:quiz_tuple[0].index('=') + 1] + ' ____',
                              quiz_tuple[1][:quiz_tuple[1].index('=') + 1] + ' ____'))
        if mask == 1:
            flag_a = random.randint(0, elements-1)
            column_a  = re.findall('[0-9]+', quiz_tuple[0])
            column_a_map = re.sub('\w', '', quiz_tuple[0]).replace(' ', '')
            flag_b = random.randint(0, elements-1)
            column_b = re.findall('[0-9]+', quiz_tuple[1])
            column_b_map = re.sub('\w', '', quiz_tuple[1]).replace(' ', '')
                ## result quiz_tuple [('17 + 13 + 8 = 38', '3 + 15 + 19 = 37')]
                ## column_a = ['17', '13', '8', '38'] column_b = ['3', '15', '19', '37']
                ## column_a_map = '++='   &           column_b_map = '++='
            column_a[flag_a] = "___"
            column_b[flag_b] = "___"
            new_quiz_column_a = ''
            new_quiz_column_b = ''
            for i in range(len(column_a)-1):
                new_quiz_column_a = new_quiz_column_a + ' ' + column_a[i] + ' ' + column_a_map[i]
                new_quiz_column_b = new_quiz_column_b + ' ' + column_b[i] + ' ' + column_b_map[i]
            new_table.append((new_quiz_column_a + ' ' + column_a[-1], new_quiz_column_b+ ' ' + column_b[-1]))

    return new_table


def create_doc(i):
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = '仿宋'
    font.size = Pt(10)
    # number = ''
    max_num = 0
    min_num = 999999
    for k in obj:
        # number += str(k) + ','
        if k > max_num:
            max_num = k
        elif k < min_num:
            min_num = k
    for k in field:
        if k > max_num:
            max_num = k
        elif k < min_num:
            min_num = k
    # number = number[:-1]

    title1 = target_date + ' -- '
    title2 = calc.upper()
    title3 = ' of integers between '
    title4 = str(min_num) + ' ~ ' + str(max_num)
    title5 = trans(calc)

    for page in range(i):
        # title creation
        p = document.add_heading(level=0)
        wp = p.add_run(title1)
        wp.font.size = Pt(15)
        wp.font.bold = True
        wp.font.color.rgb = RGBColor(0, 0, 0)

        wp = p.add_run(title2)
        wp.font.size = Pt(15)
        wp.font.color.rgb = RGBColor(255, 0, 0)

        wp = p.add_run(title3)
        wp.font.size = Pt(15)
        wp.font.bold = True
        wp.font.color.rgb = RGBColor(0, 0, 0)

        wp = p.add_run(title4 + '\n')
        wp.font.size = Pt(15)
        wp.font.color.rgb = RGBColor(255, 0, 0)

        wp = p.add_run('整数 ')
        wp.font.size = Pt(15)
        wp.font.bold = True
        wp.font.color.rgb = RGBColor(0, 0, 0)

        wp = p.add_run(title4)
        wp.font.size = Pt(15)
        wp.font.color.rgb = RGBColor(255, 0, 0)

        wp = p.add_run(' 之间的')
        wp.font.size = Pt(15)
        wp.font.bold = True
        wp.font.color.rgb = RGBColor(0, 0, 0)

        wp = p.add_run(title5)
        wp.font.size = Pt(15)
        wp.font.color.rgb = RGBColor(255, 0, 0)

        wp = p.add_run('练习 -- 第 ' + str(page + 1) + ' 页')
        wp.font.size = Pt(15)
        wp.font.bold = True
        wp.font.color.rgb = RGBColor(0, 0, 0)

        if len(field2) == 1:
            quiz_table = quiz_create(field, obj, calc, quiz_number, max_result, 2)
            # like this [('4 + 3 = 7', '3 + 2 = 5'), ('3 + 12 = 15', '7 + 4 = 11'), ('2 + 2 = 4', '2 + 5 = 7'),
            # ('13 + 3 = 16', '1 + 2 = 3'), ('11 + 2 = 13', '10 + 2 = 12'), ('14 + 3 = 17', '6 + 4 = 10'),
            # ('3 + 8 = 11', '3 + 9 = 12'), ('2 + 9 = 11', '3 + 14 = 17'), ('3 + 2 = 5', '6 + 3 = 9'),
            # ('2 + 1 = 3', '3 + 12 = 15')]
        else:
            quiz_table = quiz_create(field, obj, calc, quiz_number, max_result, 3)
            # like this [('4 + 3  + 0= 7', '3 + 2  + 0= 5'), .....]

        masked_table = mask_table(quiz_table, quiz_type[page], elements=len(field2) + 1)
        table = document.add_table(rows=1, cols=2)

        for q1, q2 in masked_table:
            row_cells = table.add_row().cells
            row_cells[0].text = q1
            row_cells[1].text = q2

        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size = Pt(quiz_font_size)
        if page < i -1:
            document.add_page_break()
    try:
        document.save('f:\\___SOPHIA____\\--Sophia K2 Folder\\MATH_PRACTICE\\math_quiz' + '_' + \
                      target_date + '.docx')
    except FileNotFoundError:
        document.save('temp.docx')


adv_day = int(input('how many days from now? (1=today)')) + 0
for k in range(adv_day):
    target_date = (datetime.now() + timedelta(days=k)).date().strftime('%Y_%m_%d')

    #obj = [10, 9, 8, 7, 8, 5, 4, 3, 2, 1]          # 2 to plus the other number
    obj = random.sample(range(10, 30), 20)
    quiz_type = [0, 0, 0]       # 0 is normal quiz, 1 is left either first or second blank in the rest 2 pages.
    field = [10, 80]     # 0 ~ 12 number to plus
    field2 = [1, ]     # if length of field 2 is less than 2, then just 2 elements added
    max_result = 100
    calc = 'addition'
    if len(field2) == 1:
        ## (32/32 @ pt16) (22/28 @ pt26)
        quiz_number = 22  # quiz per page in 2 columns  20 is the max per page and should be even number
        quiz_font_size = 26
        max_quizs_per_page = 28
    else:
        quiz_number = 22  # quiz per page in 2 columns  20 is the max per page and should be even number
        quiz_font_size = 18
        max_quizs_per_page = 28


    create_doc(len(quiz_type))   # 3 pages per quiz

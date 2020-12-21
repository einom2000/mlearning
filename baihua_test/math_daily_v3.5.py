import random
import re
from datetime import datetime, timedelta

from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor


class QuizGenerator(object):
    def __init__(self, type, max_quiz_per_page):
        self.type = type
        self.max_quiz_per_page = max_quiz_per_page
        self.elements = type.replace('?', '').find('=') + 1
        self.fill_in = False
        if type.find('?') >= 0:
            self.fill_in = True
        self.max = int(type[type.find('=') + 1:type.find('[')])
        self.range = (int(type[type.find('[') + 1:type.find(']')].split(',')[0]),
                      int(type[type.find('[') + 1:type.find(']')].split(',')[1]))

        self.quiz_per_page = []
        self.type_per_page = []

    def quiz_generate(self):
        for _ in range(self.max_quiz_per_page):
            self.quiz_per_page.append(list(self.random_quiz()))
            self.type_per_page.append(self.type[:self.type.find('=') + 1])
        # print(quiz_per_page)
        # print(type_per_page)
        # [(['66', '36', '21'], 9), (['72', '11', '58'], 3), (['90', '36', '26'], 28), ...]
        # ['--=', '--=', '--=', '--=', '--=', '--=', '--=', '--=', '--=', '--=', '--=', ...]

    def random_quiz(self):
        while True:
            elements_list = []
            random_number = str(random.randint(self.range[0], self.range[1]))
            single_quiz = random_number
            elements_list.append(random_number)
            for i in range(self.elements - 1):
                single_quiz += self.type.replace('?', '')[i]
                random_number = str(random.randint(self.range[0], self.range[1]))
                single_quiz += random_number
                elements_list.append(random_number)

            if 0 < eval(single_quiz) <= self.max:
                return elements_list, eval(single_quiz)


def generate_one_day(date, page_per_day, file_name_surfix, types_per_page, max_quiz_per_page):
    quiz_doc_file = []
    for i in range(page_per_day):
        quiz_page = QuizGenerator(types_per_page[i], max_quiz_per_page)
        quiz_doc_file.append(quiz_page)
        quiz_doc_file[i].quiz_generate()
    create_doc(date, page_per_day, file_name_surfix, quiz_doc_file)


def create_doc(date, page_per_day, file_name_surfix, quiz_doc_file):
    global quiz_font_size
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = '仿宋'
    font.size = Pt(10)

    title1 = date + ' -- '
    title3 = ' of integers between '

    for page in range(page_per_day):
        title2 = type_verbose(quiz_doc_file[page].type_per_page[0])
        title4 = str(quiz_doc_file[page].range[0]) + ' ~ ' + str(quiz_doc_file[page].range[1])
        # title creation
        p = document.add_heading(level=0)
        wp = p.add_run(title1)
        wp.font.size = Pt(15)
        wp.font.bold = True
        wp.font.color.rgb = RGBColor(0, 0, 0)

        wp = p.add_run(title2)
        wp.font.size = Pt(15)
        wp.font.color.rgb = RGBColor(255, 0, 0)

        wp = p.add_run(title3)
        wp.font.size = Pt(15)
        wp.font.bold = True
        wp.font.color.rgb = RGBColor(0, 0, 0)

        wp = p.add_run(title4 + '\n')
        wp.font.size = Pt(15)
        wp.font.color.rgb = RGBColor(255, 0, 0)

        wp = p.add_run('整数 ')
        wp.font.size = Pt(15)
        wp.font.bold = True
        wp.font.color.rgb = RGBColor(0, 0, 0)

        wp = p.add_run(title4)
        wp.font.size = Pt(15)
        wp.font.color.rgb = RGBColor(255, 0, 0)

        wp = p.add_run(' 之间的')
        wp.font.size = Pt(15)
        wp.font.bold = True
        wp.font.color.rgb = RGBColor(0, 0, 0)

        wp = p.add_run(title2)
        wp.font.size = Pt(15)
        wp.font.color.rgb = RGBColor(255, 0, 0)

        wp = p.add_run('练习 -- 第 ' + str(page + 1) + ' 页')
        wp.font.size = Pt(15)
        wp.font.bold = True
        wp.font.color.rgb = RGBColor(0, 0, 0)

        # print(quiz_doc_file[page].type_per_page)
        # print(quiz_doc_file[page].quiz_per_page)
        quiz_font_size = 22 - (len(quiz_doc_file[page].quiz_per_page[0][0]) - 2) * 3

        masked_table = get_masked_table(quiz_doc_file[page])
        table = document.add_table(rows=1, cols=2)

        for q1, q2 in masked_table:
            row_cells = table.add_row().cells
            row_cells[0].text = q1
            row_cells[1].text = q2

        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size = Pt(quiz_font_size)
        if page < page_per_day - 1:
            document.add_page_break()
    try:
        document.save('f:\\___SOPHIA____\\--Sophia K3 Folder\\MATH_PRACTICE\\math_quiz' + file_name_surfix + '_' + \
                      date + '.docx')
    except FileNotFoundError:
        document.save('temp.docx')


def get_masked_table(quiz_of_this_page):
    # [(['66', '36', '21'], 9), (['72', '11', '58'], 3), (['90', '36', '26'], 28), ...]
    # ['--=', '--=', '--=', '--=', '--=', '--=', '--=', '--=', '--=', '--=', '--=', ...]
    masked_queue = []
    if quiz_of_this_page.type_per_page[0].find('?') > 0:
        for quiz in quiz_of_this_page.quiz_per_page:
            k = random.randint(0, len(quiz[0])-1)
            quiz[0][k] = '__'
    else:
        for quiz in quiz_of_this_page.quiz_per_page:
            quiz[1] = '__'
    type = quiz_of_this_page.type_per_page[0].replace('?', '')
    for i in range(0, len(quiz_of_this_page.type_per_page), 2):
        quiz_rowed = []
        for j in range(2):
            quiz = quiz_of_this_page.quiz_per_page[i + j]
            temp = ''
            for k in range(len(quiz[0])):
                temp += quiz[0][k] + ' ' + type[k] + ' '
            temp = temp + ' ' + str(quiz[1])
            quiz_rowed.append(temp)
        masked_queue.append(quiz_rowed)
    # print(masked_queue)
    print(masked_queue)
    return masked_queue


def type_verbose(type):
    tmp = 'bcdefghijklmnopqrstuvwxyz'
    verbose = 'a'
    type1 = type.replace('?', '')
    type1 = type1.replace('=', '')
    for i in range(len(type1)):
        verbose += type1[i]
        verbose += tmp[i]
    if type.find('?') > 0:
        temp = verbose
        verbose = temp[:type.find('?') + 1] + '?' + temp[type.find('?') + 2:]
    return verbose + '=' + tmp[len(type1)]


def get_genral_input():
    # get the dates to print
    try:
        adv_date = int(input('How many days from now? (default = 1(today), max = 30days)')) + 0
        if adv_date > 30 or adv_date <= 0:
            adv_date = 1
    except ValueError:
        adv_date = 1

    # get how many pages per day:
    try:
        page_per_day = int(input('How many pages per day? (default = 3 pages, max = 5 pages)'))
        if page_per_day >5 or page_per_day <= 0:
            page_per_day = 1
    except ValueError:
        page_per_day = 3

    try:
        # get max quizes per page:
        max_quiz_per_page = int(input('Pleas input max quiz per page? (default = 22 quiz)')) // 2 * 2
        if max_quiz_per_page > 22 or max_quiz_per_page <= 10:
            max_quiz_per_page = 22
    except ValueError:
        max_quiz_per_page = 22
    print(adv_date, page_per_day, max_quiz_per_page)

    # get file_name_surfix
    file_name_surfix = input("file name surfix:")
    if file_name_surfix != '':
        file_name_surfix = '_' + file_name_surfix

    return adv_date, page_per_day, file_name_surfix, max_quiz_per_page


def get_quiz_type_for_each_page(page):
    print('Please enter the %d page\'s quiz type' % page)
    print('type \'a+b+c+d=100[10, 99]\' stands for abcd all in [10,99] range and the result <=100')
    type = input('please give a type: (use enter to use previous or default(a+b=100)')
    if type == '' and page != 1:
        return 'previous'
    elif (page == 1 and type == '') or type.find('[') < 0 or type.find(']') < 0 or type.find(',') < 0 \
            or type.find('=') < 0:
        print('Use the default a+b=100[10,99]')
        return '+=100[10,99]'
    else:
        type1 = re.sub('[a-zA-Z!"#$%&\'().:;@\\^_`{|}~ \t\n\r\x0b\x0c]', '', type).replace(' ', '')
    return type1


def main():
    global quiz_font_size
    quiz_font_size = 22
    adv_date, page_per_day, file_name_surfix, max_quiz_per_page = get_genral_input()
    types_per_page = []
    for i in range(page_per_day):
        temp = get_quiz_type_for_each_page(i + 1)
        if temp != 'previous':
            types_per_page.append(temp)
        else:
            types_per_page.append(types_per_page[i-1])
    today = 'sample'
    for i in range(adv_date):
        date = (datetime.now() + timedelta(days=i)).date().strftime('%Y_%m_%d')
        generate_one_day(date, page_per_day, file_name_surfix, types_per_page, max_quiz_per_page)
    pass


if __name__ == '__main__':
    main()

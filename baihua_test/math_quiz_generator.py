import random
from datetime import datetime, timedelta

from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor


# translation function
def trans(calc):
    tmp = ''
    if calc == 'division':
        tmp = '除法'
    elif calc == 'subtraction':
        tmp = '减法'
    elif calc == 'multiplication':
        tmp = '乘法'
    else:
        tmp = '加法'
    return tmp


# create quiz list
def quiz_create(field, obj, calc, quiz_number, max_result):
    calc_dic = {'addition': '+',
                'subtraction': '-',
                'multiplication': 'x',
                'division': '÷'}
    length = quiz_number // 2 * 2
    subjs = []
    for _ in range(0, length):
        subjs.append(random.randint(field[0], field[1]))
    quizs = []
    for subj in subjs:
        while True:
            tmp = [subj, random.choice(obj)]
            chk_result = eval(str(tmp[0]) + calc_dic[calc] + str(tmp[1]))
            if tmp not in quizs and [tmp[1], tmp[0]] not in quizs and [tmp[0], tmp[1]] not in quizs \
                    and chk_result <= max_result:
                break
        random.shuffle(tmp)
        quizs.append(tmp)
    quiz_table = []
    quiz_table1 = []
    for i in range(0, len(quizs), 2):
        quiz_table.append((str(quizs[i][0]) + ' ' + calc_dic[calc] + ' ' + str(quizs[i][1]) + ' =',
                           str(quizs[i + 1][0]) + ' ' + calc_dic[calc] + ' ' + str(quizs[i + 1][1]) + ' ='))
    ready_quiz_table = make_up_final_quiz(quiz_table + quiz_table1)
    return ready_quiz_table


def make_up_final_quiz(table):
    new_table = []
    # put actual result to the table
    for quiz_tuple in table:
        new_table.append((quiz_tuple[0] + ' ' + str(eval(quiz_tuple[0].replace('=', ''))),
                          quiz_tuple[1] + ' ' + str(eval(quiz_tuple[1].replace('=', '')))))
    return new_table


def mask_table(table, mask):
    new_table = []
    calc_dic = {'addition': '+',
                'subtraction': '-',
                'multiplication': 'x',
                'division': '÷'}
    for quiz_tuple in table:
        if mask == 0:
            new_table.append((quiz_tuple[0][:quiz_tuple[0].index('=') + 1] + ' ____',
                              quiz_tuple[1][:quiz_tuple[1].index('=') + 1] + ' ____'))
        if mask == 1:
            flag = random.choice([0, 1])
            if not flag:
                new_table.append(("____ " + quiz_tuple[0][quiz_tuple[0].index(calc_dic[calc]):],
                                  "____ " + quiz_tuple[1][quiz_tuple[1].index(calc_dic[calc]):]))
            else:
                new_table.append((quiz_tuple[0][:quiz_tuple[0].index(calc_dic[calc]) + 1] + ' ____ ' +
                                  quiz_tuple[0][quiz_tuple[0].index('='):],
                                  quiz_tuple[1][:quiz_tuple[1].index(calc_dic[calc]) + 1] + ' ____ ' +
                                  quiz_tuple[1][quiz_tuple[1].index('='):]))
    return new_table


def create_doc(i):
    document = Document()
    number = ''
    for k in obj:
        number += str(k) + ','
    number = number[:-1]

    title1 = target_date + ' -- '
    title2 = calc.upper()
    title3 = ' of number '
    title4 = str(number) + '\n'
    title5 = trans(calc)

    for page in range(i):
        # title creation
        p = document.add_heading(level=0)
        wp = p.add_run(title1)
        wp.font.size = Pt(18)
        wp.font.bold = True
        wp.font.color.rgb = RGBColor(0, 0, 0)
        wp = p.add_run(title2)
        wp.font.size = Pt(18)
        wp.font.color.rgb = RGBColor(255, 0, 0)
        wp = p.add_run(title3)
        wp.font.size = Pt(18)
        wp.font.bold = True
        wp.font.color.rgb = RGBColor(0, 0, 0)
        wp = p.add_run(title4)
        wp.font.size = Pt(18)
        wp.font.color.rgb = RGBColor(255, 0, 0)

        wp = p.add_run('关于数字 ')
        wp.font.name = '#simSong'
        wp.font.size = Pt(18)
        wp.font.bold = True
        wp.font.color.rgb = RGBColor(0, 0, 0)
        wp = p.add_run(str(number))
        wp.font.size = Pt(18)
        wp.font.color.rgb = RGBColor(255, 0, 0)
        wp = p.add_run(' 的')
        wp.font.size = Pt(18)
        wp.font.bold = True
        wp.font.color.rgb = RGBColor(0, 0, 0)
        wp = p.add_run(title5)
        wp.font.size = Pt(18)
        wp.font.color.rgb = RGBColor(255, 0, 0)
        wp = p.add_run('练习 -- 第 ' + str(page + 1) + ' 页')
        wp.font.size = Pt(18)
        wp.font.bold = True
        wp.font.color.rgb = RGBColor(0, 0, 0)

        quiz_table = quiz_create(field, obj, calc, quiz_number, max_result)
        # like this [('4 + 3 = 7', '3 + 2 = 5'), ('3 + 12 = 15', '7 + 4 = 11'), ('2 + 2 = 4', '2 + 5 = 7'),
        # ('13 + 3 = 16', '1 + 2 = 3'), ('11 + 2 = 13', '10 + 2 = 12'), ('14 + 3 = 17', '6 + 4 = 10'),
        # ('3 + 8 = 11', '3 + 9 = 12'), ('2 + 9 = 11', '3 + 14 = 17'), ('3 + 2 = 5', '6 + 3 = 9'),
        # ('2 + 1 = 3', '3 + 12 = 15')]
        masked_table = mask_table(quiz_table, quiz_type[page])
        table = document.add_table(rows=1, cols=2)

        for q1, q2 in masked_table:
            row_cells = table.add_row().cells
            row_cells[0].text = q1
            row_cells[1].text = q2

        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size = Pt(26)
        if page < i -1:
            document.add_page_break()
    try:
        document.save('f:\\___SOPHIA____\\--Sophia K2 Folder\\MATH_PRACTICE\\math_quiz' + '_' + \
                      target_date + '.docx')
    except FileNotFoundError:
        document.save('temp.docx')


adv_day = int(input('how many days from now? (1=today)'))
for k in range(adv_day):
    target_date = (datetime.now() + timedelta(days=k)).date().strftime('%Y_%m_%d')

    obj = [10, 9, 8, 7, 8, 5, 4, 3, 2, 1]          # 2 to plus the other number
    quiz_type = [0, 0, 1]       # 0 is normal quiz, 1 is left either first or second blank in the rest 2 pages.
    quiz_number = 22   # quiz per page in 2 columns  20 is the max per page and should be even number
    field = [1, 20]     # 0 ~ 12 number to plus
    max_result = 20
    calc = 'addition'
    max_quizs_per_page = 28

    create_doc(len(quiz_type))   # 3 pages per quiz
